import io
from dataclasses import dataclass, field
from typing import Optional

from PIL import Image
from docx import Document


@dataclass
class ParsedDocument:
    filename: str
    text: str = ""
    images: list[Image.Image] = field(default_factory=list)
    pdf_bytes: Optional[bytes] = None
    mime_type: Optional[str] = None
    error: Optional[str] = None


def _extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract readable text from a DOCX file.

    Includes:
    - normal paragraphs
    - headings
    - table contents
    """

    document = Document(io.BytesIO(file_bytes))

    parts: list[str] = []

    # ---------------------------------------------------------
    # Paragraphs
    # ---------------------------------------------------------
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if not text:
            continue

        style_name = (
            paragraph.style.name
            if paragraph.style is not None
            else ""
        )

        # Preserve headings in a Markdown-like form
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
                level = min(max(level, 1), 6)
                parts.append(f"{'#' * level} {text}")
            except Exception:
                parts.append(text)
        else:
            parts.append(text)

    # ---------------------------------------------------------
    # Tables
    # ---------------------------------------------------------
    for table_index, table in enumerate(document.tables, start=1):

        rows = []

        for row in table.rows:
            cells = [
                cell.text.strip().replace("\n", " ")
                for cell in row.cells
            ]

            if any(cells):
                rows.append(cells)

        if not rows:
            continue

        parts.append(f"\nTable {table_index}:")

        for row in rows:
            parts.append(
                " | ".join(cell or "" for cell in row)
            )

    return "\n".join(parts).strip()


def extract_content_from_file(
    filename: str,
    file_bytes: bytes,
) -> ParsedDocument:

    """
    Validate and extract content from supported uploads.

    Supported:
    - PDF
    - PNG / JPG / JPEG / WEBP
    - DOCX
    - TXT / MD / CSV
    """

    result = ParsedDocument(filename=filename)

    lower_name = filename.lower()

    if not file_bytes:
        result.error = "The uploaded file is empty."
        return result

    # =========================================================
    # PDF
    # =========================================================
    if lower_name.endswith(".pdf"):

        if len(file_bytes) > 50 * 1024 * 1024:
            result.error = (
                "The PDF exceeds the 50 MB processing limit."
            )
            return result

        # Preserve PDF directly for Gemini multimodal processing
        result.pdf_bytes = file_bytes
        result.mime_type = "application/pdf"

        return result

    # =========================================================
    # IMAGES
    # =========================================================
    if lower_name.endswith(
        (".png", ".jpg", ".jpeg", ".webp")
    ):

        try:
            image = Image.open(
                io.BytesIO(file_bytes)
            ).convert("RGB")

            result.images.append(image)

            result.mime_type = (
                Image.MIME.get(
                    Image.open(io.BytesIO(file_bytes)).format,
                    "image/jpeg",
                )
            )

        except Exception as exc:
            result.error = (
                f"Could not open the image: {exc}"
            )

        return result

    # =========================================================
    # DOCX
    # =========================================================
    if lower_name.endswith(".docx"):

        try:
            text = _extract_docx_text(file_bytes)

            if not text:
                result.error = (
                    "The DOCX file contains no readable text."
                )
                return result

            result.text = text

            result.mime_type = (
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            )

        except Exception as exc:
            result.error = (
                f"Could not read the DOCX file: {exc}"
            )

        return result

    # =========================================================
    # TEXT FILES
    # =========================================================
    if lower_name.endswith(
        (".txt", ".md", ".csv")
    ):

        try:
            result.text = file_bytes.decode(
                "utf-8",
                errors="ignore",
            ).strip()

            if not result.text:
                result.error = (
                    "The uploaded text file contains "
                    "no readable content."
                )
                return result

            if lower_name.endswith(".md"):
                result.mime_type = "text/markdown"

            elif lower_name.endswith(".csv"):
                result.mime_type = "text/csv"

            else:
                result.mime_type = "text/plain"

        except Exception as exc:
            result.error = (
                f"Could not decode the text file: {exc}"
            )

        return result

    # =========================================================
    # UNSUPPORTED
    # =========================================================
    result.error = (
        "Unsupported file type. "
        "Upload a PDF, DOCX, image, TXT, Markdown, or CSV file."
    )

    return result